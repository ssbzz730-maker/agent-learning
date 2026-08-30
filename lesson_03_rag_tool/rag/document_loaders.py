"""把 TXT、Markdown、PDF 和 DOCX 转换成统一文本段落。"""

import re
from pathlib import Path


SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def clean_text(text):
    """保留换行结构，并清理重复空格和空行。"""

    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def make_section(text, source, file_type, **extra_metadata):
    """生成所有加载器共用的 section 结构。"""

    metadata = {
        "source": source,
        "file_name": Path(source).name,
        "file_type": file_type,
        **extra_metadata,
    }
    return {"text": clean_text(text), "metadata": metadata}


def load_txt(path, source=None):
    source = source or path.name
    section = make_section(path.read_text(encoding="utf-8-sig"), source, "txt")
    return [section] if section["text"] else []


def load_markdown(path, source=None):
    """按 Markdown 标题分段，并把标题写入 metadata。"""

    source = source or path.name
    sections = []
    heading = "文档开头"
    body = []

    def save_section():
        section = make_section("\n".join(body), source, "md", heading=heading)
        if section["text"]:
            sections.append(section)

    for line in path.read_text(encoding="utf-8-sig").splitlines():
        match = re.match(r"^#{1,6}\s+(.+?)\s*$", line)
        if match:
            save_section()
            heading = match.group(1)
            body = []
        else:
            body.append(line)
    save_section()
    return sections


def load_pdf(path, source=None):
    """逐页读取 PDF；依赖只在真正读取 PDF 时加载。"""

    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError("读取 PDF 需要安装 pypdf") from error

    source = source or path.name
    sections = []
    for page_number, page in enumerate(PdfReader(str(path)).pages, start=1):
        section = make_section(
            page.extract_text() or "",
            source,
            "pdf",
            page=page_number,
        )
        if section["text"]:
            sections.append(section)
    return sections


def load_docx(path, source=None):
    """读取 Word 段落和表格，标题样式保存为 heading。"""

    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("读取 DOCX 需要安装 python-docx") from error

    source = source or path.name
    document = Document(str(path))
    sections = []
    heading = "文档开头"
    body = []

    def save_section():
        section = make_section("\n".join(body), source, "docx", heading=heading)
        if section["text"]:
            sections.append(section)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            save_section()
            heading = text
            body = []
        else:
            body.append(text)

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            body.append("\n".join(rows))
    save_section()
    return sections


def load_file(path, document_dir=None):
    """根据扩展名调用加载器，返回统一 section 列表。"""

    path = Path(path)
    source = (
        path.relative_to(document_dir).as_posix()
        if document_dir is not None
        else path.name
    )
    loaders = {
        ".txt": load_txt,
        ".md": load_markdown,
        ".markdown": load_markdown,
        ".pdf": load_pdf,
        ".docx": load_docx,
    }
    suffix = path.suffix.lower()
    if suffix not in loaders:
        raise ValueError(f"不支持的文件格式：{suffix}")
    return loaders[suffix](path, source)
